import os
import PyPDF2
import docx
from typing import Optional
import aiofiles

class ResumeProcessor:
    """Сервис для обработки и извлечения текста из резюме различных форматов"""
    
    async def extract_text(self, file_path: str) -> str:
        """Извлечение текста из файла резюме"""
        
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"Файл {file_path} не найден")
        
        file_extension = os.path.splitext(file_path)[1].lower()
        
        try:
            if file_extension == '.pdf':
                return await self._extract_from_pdf(file_path)
            elif file_extension in ['.doc', '.docx']:
                return await self._extract_from_docx(file_path)
            elif file_extension == '.txt':
                return await self._extract_from_txt(file_path)
            else:
                raise ValueError(f"Неподдерживаемый формат файла: {file_extension}")
                
        except Exception as e:
            raise Exception(f"Ошибка извлечения текста: {str(e)}")
    
    async def _extract_from_pdf(self, file_path: str) -> str:
        """Извлечение текста из PDF файла"""
        
        text = ""
        
        try:
            with open(file_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                
                for page_num in range(len(pdf_reader.pages)):
                    page = pdf_reader.pages[page_num]
                    text += page.extract_text() + "\n"
                    
        except Exception as e:
            raise Exception(f"Ошибка чтения PDF: {str(e)}")
        
        return self._clean_text(text)
    
    async def _extract_from_docx(self, file_path: str) -> str:
        """Извлечение текста из DOCX файла"""
        
        try:
            doc = docx.Document(file_path)
            text = ""
            
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
                
            # Также извлекаем текст из таблиц
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
                    
        except Exception as e:
            raise Exception(f"Ошибка чтения DOCX: {str(e)}")
        
        return self._clean_text(text)
    
    async def _extract_from_txt(self, file_path: str) -> str:
        """Извлечение текста из TXT файла"""
        
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8') as file:
                text = await file.read()
        except UnicodeDecodeError:
            # Пробуем другие кодировки
            try:
                async with aiofiles.open(file_path, 'r', encoding='cp1251') as file:
                    text = await file.read()
            except:
                async with aiofiles.open(file_path, 'r', encoding='latin-1') as file:
                    text = await file.read()
        except Exception as e:
            raise Exception(f"Ошибка чтения TXT: {str(e)}")
        
        return self._clean_text(text)
    
    def _clean_text(self, text: str) -> str:
        """Очистка и нормализация текста"""
        
        if not text:
            return ""
        
        # Убираем лишние пробелы и переносы строк
        lines = text.split('\n')
        cleaned_lines = []
        
        for line in lines:
            line = line.strip()
            if line:  # Пропускаем пустые строки
                cleaned_lines.append(line)
        
        # Объединяем строки
        cleaned_text = '\n'.join(cleaned_lines)
        
        # Убираем множественные пробелы
        import re
        cleaned_text = re.sub(r'\s+', ' ', cleaned_text)
        
        return cleaned_text.strip()
    
    def validate_resume_content(self, text: str) -> bool:
        """Проверка, что текст похож на резюме"""
        
        if not text or len(text) < 100:
            return False
        
        # Ключевые слова, которые обычно встречаются в резюме
        resume_keywords = [
            'опыт', 'работа', 'образование', 'навыки', 'достижения',
            'проект', 'компания', 'должность', 'технологии', 'languages',
            'experience', 'education', 'skills', 'achievements', 'position',
            'university', 'degree', 'certificate', 'курсы', 'сертификат'
        ]
        
        text_lower = text.lower()
        keyword_count = sum(1 for keyword in resume_keywords if keyword in text_lower)
        
        # Если найдено хотя бы 3 ключевых слова, считаем что это резюме
        return keyword_count >= 3
    
    def extract_contact_info(self, text: str) -> dict:
        """Извлечение контактной информации из резюме"""
        
        import re
        
        contact_info = {
            'emails': [],
            'phones': [],
            'links': []
        }
        
        # Поиск email адресов
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        emails = re.findall(email_pattern, text)
        contact_info['emails'] = list(set(emails))
        
        # Поиск телефонов (российские номера)
        phone_patterns = [
            r'\+7[\s\-\(\)]?\d{3}[\s\-\(\)]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}',
            r'8[\s\-\(\)]?\d{3}[\s\-\(\)]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}',
            r'\d{3}[\s\-]?\d{3}[\s\-]?\d{2}[\s\-]?\d{2}'
        ]
        
        phones = []
        for pattern in phone_patterns:
            phones.extend(re.findall(pattern, text))
        contact_info['phones'] = list(set(phones))
        
        # Поиск ссылок (LinkedIn, GitHub и т.д.)
        link_pattern = r'https?://[^\s<>"{}|\\^`\[\]]+'
        links = re.findall(link_pattern, text)
        contact_info['links'] = list(set(links))
        
        return contact_info